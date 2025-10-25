import os
import platform
import logging
import tempfile

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def convert_docx_to_pdf(input_path: str, output_filename: str = None) -> str:
    if output_filename:
        output_path = os.path.join(os.path.dirname(input_path), output_filename)
    else:
        output_path = os.path.splitext(input_path)[0] + ".pdf"
    
    # Try using docx2pdf on Windows first (uses Microsoft Word COM - best quality & Unicode support)
    if platform.system() == "Windows":
        try:
            from docx2pdf import convert
            logger.info("Attempting conversion with docx2pdf (MS Word)...")
            
            abs_input = os.path.abspath(input_path)
            abs_output = os.path.abspath(output_path)
            
            logger.info(f"Converting: {abs_input}")
            logger.info(f"Output to: {abs_output}")
            
            # Remove output file if it exists
            if os.path.exists(abs_output):
                os.remove(abs_output)
            
            convert(abs_input, abs_output)
            
            # Wait a moment for file to be written
            import time
            time.sleep(0.5)
            
            if os.path.exists(abs_output) and os.path.getsize(abs_output) > 0:
                logger.info(f"✓ Successfully converted using docx2pdf (MS Word) - Size: {os.path.getsize(abs_output)} bytes")
                return abs_output
            else:
                raise Exception("Output PDF was not created or is empty")
                
        except Exception as e:
            logger.error(f"docx2pdf failed: {str(e)}")
            
            # Try direct Word COM automation as fallback
            try:
                logger.info("Trying direct Word COM automation...")
                import win32com.client
                import time
                
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                abs_input = os.path.abspath(input_path)
                abs_output = os.path.abspath(output_path)
                
                logger.info(f"Opening document: {abs_input}")
                
                # Remove output if exists
                if os.path.exists(abs_output):
                    os.remove(abs_output)
                
                # Open document
                doc = word.Documents.Open(abs_input)
                
                # Save as PDF (FileFormat 17 = wdFormatPDF)
                logger.info(f"Saving as PDF: {abs_output}")
                doc.SaveAs(abs_output, FileFormat=17)
                
                # Close document
                doc.Close(False)
                word.Quit()
                
                # Wait for file to be completely written
                time.sleep(1)
                
                if os.path.exists(abs_output) and os.path.getsize(abs_output) > 0:
                    logger.info(f"✓ Successfully converted using Word COM - Size: {os.path.getsize(abs_output)} bytes")
                    return abs_output
                else:
                    raise Exception("Word COM failed to create PDF")
                    
            except Exception as e2:
                logger.error(f"Word COM also failed: {str(e2)}")
                try:
                    # Cleanup Word in case of error
                    word.Quit()
                except:
                    pass
                logger.info("Falling back to alternative methods...")
    
    # Try pypandoc direct PDF conversion first (best for math formulas and Unicode)
    try:
        import pypandoc
        logger.info("Attempting direct DOCX to PDF conversion with pypandoc...")
        
        # Try with XeLaTeX (best Unicode and math support)
        try:
            pypandoc.convert_file(
                input_path,
                'pdf',
                outputfile=output_path,
                extra_args=[
                    '--pdf-engine=xelatex',
                    '--variable', 'mainfont:Nirmala UI',
                    '--variable', 'CJKmainfont:SimSun',
                    '--variable', 'mathfont:Latin Modern Math',
                    '-V', 'geometry:margin=1in',
                ]
            )
            
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                logger.info("✓ Successfully converted with pypandoc+XeLaTeX")
                return output_path
                
        except RuntimeError as e:
            if 'xelatex not found' in str(e).lower():
                logger.warning("XeLaTeX not available, trying HTML method...")
            else:
                logger.warning(f"pypandoc direct conversion failed: {e}")
    except Exception as e:
        logger.warning(f"pypandoc not available or failed: {e}")
    
    # Try pypandoc for HTML conversion (better for math and Unicode)
    try:
        logger.info("Converting DOCX → HTML → PDF with pypandoc...")
        
        # Step 1: Convert DOCX to HTML using pypandoc (better than mammoth for math/Unicode)
        import pypandoc
        
        html_temp_path = output_path.replace('.pdf', '_temp.html')
        
        try:
            # Convert DOCX to HTML with MathML support
            pypandoc.convert_file(
                input_path,
                'html',
                outputfile=html_temp_path,
                extra_args=[
                    '--mathml',  # Convert math to MathML
                    '--standalone',  # Create complete HTML
                    '--self-contained',  # Embed all resources
                ]
            )
            logger.info("✓ Converted DOCX to HTML using pypandoc")
            html_content = None  # Will read from file
            
        except Exception as e:
            logger.warning(f"pypandoc HTML conversion failed: {e}, trying mammoth...")
            
            # Fallback to mammoth
            import mammoth
            with open(input_path, "rb") as docx_file:
                result = mammoth.convert_to_html(
                    docx_file,
                    convert_image=mammoth.images.inline(lambda img: None)  # Keep images inline
                )
                html_content = result.value
                messages = result.messages
                
                if messages:
                    for msg in messages:
                        logger.debug(f"Mammoth: {msg}")
        
        # Step 2: Enhance HTML with proper CSS for Unicode and math support
        if html_content:
            # mammoth output - need to wrap in HTML structure
            enhanced_html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <style>
        @page {{
            size: A4;
            margin: 2cm;
        }}
        
        body {{
            font-family: 'Nirmala UI', 'Mangal', 'Kokila', 'Segoe UI', 'Arial', 'Tahoma', 'SimSun', sans-serif;
            font-size: 11pt;
            line-height: 1.6;
            color: #000;
            max-width: 100%;
        }}
        
        /* Explicit Hindi/Devanagari support */
        * {{
            font-family: 'Nirmala UI', 'Mangal', 'Kokila', 'Segoe UI', 'Arial', sans-serif;
        }}
        
        /* Support for RTL languages like Arabic/Urdu */
        [dir="rtl"] {{
            direction: rtl;
            text-align: right;
            font-family: 'Tahoma', 'Arabic Typesetting', 'Traditional Arabic', sans-serif;
        }}
        
        /* Math support */
        .math, math {{
            font-family: 'Cambria Math', 'STIX Two Math', 'Latin Modern Math', serif;
        }}
        
        /* Table styling */
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 10px 0;
        }}
        
        table td, table th {{
            border: 1px solid #000;
            padding: 8px;
            text-align: left;
            vertical-align: top;
        }}
        
        table th {{
            background-color: #4472C4;
            color: white;
            font-weight: bold;
        }}
        
        table tr:nth-child(even) {{
            background-color: #f2f2f2;
        }}
        
        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 1em;
            margin-bottom: 0.5em;
        }}
        
        p {{
            margin: 0.5em 0;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
            with open(html_temp_path, 'w', encoding='utf-8') as f:
                f.write(enhanced_html)
        else:
            # pypandoc output - already complete HTML, just enhance CSS
            with open(html_temp_path, 'r', encoding='utf-8') as f:
                pypandoc_html = f.read()
            
            # Add additional CSS for better rendering with Hindi font priority
            css_injection = """
    <style>
        * { font-family: 'Nirmala UI', 'Mangal', 'Kokila', 'Segoe UI', 'Arial', 'Tahoma', 'SimSun', sans-serif !important; }
        body { font-family: 'Nirmala UI', 'Mangal', 'Kokila', 'Segoe UI', 'Arial', 'Tahoma', 'SimSun', sans-serif; }
        table { border-collapse: collapse; width: 100%; margin: 10px 0; }
        table td, table th { border: 1px solid #000; padding: 8px; }
        table th { background-color: #4472C4; color: white; }
        table tr:nth-child(even) { background-color: #f2f2f2; }
        .math, math { font-family: 'Cambria Math', 'Latin Modern Math', serif; }
    </style>
</head>"""
            
            pypandoc_html = pypandoc_html.replace('</head>', css_injection)
            
            with open(html_temp_path, 'w', encoding='utf-8') as f:
                f.write(pypandoc_html)
        
        logger.info(f"✓ HTML ready: {html_temp_path}")
        
        # Step 3: Convert HTML to PDF using wkhtmltopdf (best for Unicode and math)
        try:
            import pdfkit
            logger.info("Converting HTML to PDF with wkhtmltopdf/pdfkit...")
            
            # Configure pdfkit options for better rendering with reduced margins
            options = {
                'encoding': 'UTF-8',
                'page-size': 'A4',
                'margin-top': '10mm',      # Reduced from 20mm
                'margin-right': '10mm',    # Reduced from 20mm
                'margin-bottom': '10mm',   # Reduced from 20mm
                'margin-left': '10mm',     # Reduced from 20mm
                'enable-local-file-access': None,
                'no-stop-slow-scripts': None,
                'enable-javascript': None,
                'javascript-delay': 1000,
                'dpi': 96,
                'image-quality': 94,
                'minimum-font-size': 10,
            }
            
            # Try to convert
            try:
                logger.info(f"Using margins: 10mm on all sides")
                pdfkit.from_file(html_temp_path, output_path, options=options)
                
                # Clean up temp HTML file
                if os.path.exists(html_temp_path):
                    os.remove(html_temp_path)
                
                if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                    logger.info(f"✓ Successfully converted with wkhtmltopdf - Size: {os.path.getsize(output_path)} bytes")
                    return output_path
                else:
                    raise Exception("PDF was not created")
                    
            except OSError as e:
                if 'No wkhtmltopdf executable found' in str(e):
                    logger.warning("wkhtmltopdf not found in PATH, trying xhtml2pdf...")
                    raise Exception("wkhtmltopdf not found")
                else:
                    raise
                    
        except Exception as e:
            logger.warning(f"pdfkit/wkhtmltopdf failed: {e}")
            
            # Fallback to xhtml2pdf
            try:
                from xhtml2pdf import pisa
                logger.info("Converting HTML to PDF with xhtml2pdf...")
                
                with open(html_temp_path, 'r', encoding='utf-8') as html_file:
                    html_string = html_file.read()
                
                with open(output_path, 'wb') as pdf_file:
                    pisa_status = pisa.CreatePDF(
                        html_string,
                        dest=pdf_file,
                        encoding='utf-8'
                    )
                
                # Clean up temp HTML file
                if os.path.exists(html_temp_path):
                    os.remove(html_temp_path)
                
                if not pisa_status.err and os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                    logger.info("✓ Successfully converted with xhtml2pdf")
                    return output_path
                else:
                    raise Exception(f"PDF creation had errors: {pisa_status.err}")
                    
            except Exception as e2:
                logger.warning(f"xhtml2pdf also failed: {e2}")
                # Clean up temp file
                if os.path.exists(html_temp_path):
                    os.remove(html_temp_path)
    
    except Exception as e:
        logger.error(f"HTML conversion method failed: {e}")
    
    # Fallback: Manual conversion using python-docx + reportlab with improved Unicode and table support
    logger.info("Using fallback method: python-docx + reportlab...")
    from docx import Document
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    
    doc = Document(input_path)
    pdf = SimpleDocTemplate(output_path, pagesize=letter)
    
    # Register Unicode fonts - try multiple options including multi-language support
    font_registered = False
    hindi_font_registered = False
    arabic_font_registered = False
    chinese_font_registered = False
    
    if platform.system() == "Windows":
        # Multi-language font options
        font_mappings = {
            'hindi': [
                ('C:/Windows/Fonts/Nirmala.ttf', 'NirmalaFont'),
                ('C:/Windows/Fonts/mangal.ttf', 'MangalFont'),
                ('C:/Windows/Fonts/NotoSansDevanagari-Regular.ttf', 'NotoDevanagariFont'),
            ],
            'arabic': [
                ('C:/Windows/Fonts/tahoma.ttf', 'TahomaFont'),
                ('C:/Windows/Fonts/TraditionalArabic.ttf', 'TraditionalArabicFont'),
                ('C:/Windows/Fonts/ArabicTypesetting.ttf', 'ArabicTypesettingFont'),
            ],
            'chinese': [
                ('C:/Windows/Fonts/simsun.ttc', 'SimSunFont'),
                ('C:/Windows/Fonts/msyh.ttc', 'YaHeiFont'),
                ('C:/Windows/Fonts/SimsunExtG.ttf', 'SimSunExtFont'),
                ('C:/Windows/Fonts/simsunb.ttf', 'SimSunBFont'),
            ],
            'universal': [
                ('C:/Windows/Fonts/arial.ttf', 'ArialFont'),
                ('C:/Windows/Fonts/segoeui.ttf', 'SegoeFont'),
                ('C:/Windows/Fonts/calibri.ttf', 'CalibriFont'),
            ]
        }
        
        registered_fonts = {}
        
        # Try Hindi fonts
        for font_path, font_name in font_mappings['hindi']:
            try:
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                logger.info(f"✓ Registered Hindi font: {font_name}")
                registered_fonts['hindi'] = font_name
                hindi_font_registered = True
                if not font_registered:
                    main_font = font_name
                    font_registered = True
                break
            except:
                pass
        
        # Try Arabic fonts
        for font_path, font_name in font_mappings['arabic']:
            try:
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                logger.info(f"✓ Registered Arabic/Urdu font: {font_name}")
                registered_fonts['arabic'] = font_name
                arabic_font_registered = True
                if not font_registered:
                    main_font = font_name
                    font_registered = True
                break
            except:
                pass
        
        # Try Chinese fonts
        for font_path, font_name in font_mappings['chinese']:
            try:
                # TTC files need special handling
                if font_path.endswith('.ttc'):
                    continue  # Skip TTC files for now as they need subfontIndex
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                logger.info(f"✓ Registered Chinese font: {font_name}")
                registered_fonts['chinese'] = font_name
                chinese_font_registered = True
                if not font_registered:
                    main_font = font_name
                    font_registered = True
                break
            except:
                pass
        
        # Try universal fonts as fallback
        if not font_registered:
            for font_path, font_name in font_mappings['universal']:
                try:
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    logger.info(f"Registered universal font: {font_name}")
                    registered_fonts['universal'] = font_name
                    main_font = font_name
                    font_registered = True
                    break
                except:
                    pass
        
        # Log font support status
        if registered_fonts:
            logger.info(f"Font support: {', '.join(registered_fonts.keys())}")
    
    if not font_registered:
        main_font = 'Helvetica'
        logger.warning("Using default Helvetica font - Unicode characters may not display correctly")
    
    styles = getSampleStyleSheet()
    
    # Create custom styles with Unicode font
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontName=main_font,
        fontSize=11,
        leading=16,
        alignment=TA_LEFT,
        spaceAfter=6,
        wordWrap='LTR'
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading1'],
        fontName=main_font,
        fontSize=16,
        leading=22,
        alignment=TA_LEFT,
        spaceAfter=12,
        spaceBefore=12,
        textColor=colors.HexColor('#000000'),
        bold=True
    )
    
    story = []
    
    # Add paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        text_content = para.text.strip()
        if text_content:
            # Determine style based on formatting
            if para.style.name.startswith('Heading'):
                style = heading_style
            else:
                style = body_style
            
            # Escape special XML characters but preserve Unicode
            text = text_content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Detect script type in the text
            has_devanagari = any('\u0900' <= char <= '\u097F' for char in text_content)  # Hindi
            has_arabic = any('\u0600' <= char <= '\u06FF' or '\u0750' <= char <= '\u077F' for char in text_content)  # Arabic/Urdu
            has_chinese = any('\u4E00' <= char <= '\u9FFF' or '\u3400' <= char <= '\u4DBF' for char in text_content)  # Chinese
            
            # Warn if script detected but font not available
            if has_devanagari and not hindi_font_registered:
                logger.warning(f"Paragraph {para_idx} contains Hindi text but no Hindi font available")
            if has_arabic and not arabic_font_registered:
                logger.warning(f"Paragraph {para_idx} contains Arabic/Urdu text but no Arabic font available")
            if has_chinese and not chinese_font_registered:
                logger.warning(f"Paragraph {para_idx} contains Chinese text but no Chinese font available")
            
            try:
                p = Paragraph(text, style)
                story.append(p)
                story.append(Spacer(1, 0.15 * inch))
            except Exception as e:
                logger.warning(f"Could not add paragraph {para_idx}: {e}")
                # Try encoding the text differently
                try:
                    # Create a simpler paragraph without complex formatting
                    simple_style = ParagraphStyle(
                        'SimpleStyle',
                        parent=body_style,
                        fontName=main_font,
                        fontSize=10,
                        leading=14
                    )
                    p = Paragraph(text, simple_style)
                    story.append(p)
                    story.append(Spacer(1, 0.15 * inch))
                except Exception as e2:
                    logger.error(f"Failed to add paragraph even with simple style: {e2}")
                    # Last resort: add as plain ASCII
                    safe_text = text_content.encode('ascii', 'ignore').decode('ascii')
                    if safe_text:
                        try:
                            p = Paragraph(safe_text + " [Some characters could not be displayed]", body_style)
                            story.append(p)
                            story.append(Spacer(1, 0.15 * inch))
                        except:
                            pass
    
    # Add tables with proper formatting
    for table_idx, table in enumerate(doc.tables):
        try:
            table_data = []
            max_cols = 0
            
            # First pass: determine max columns
            for row in table.rows:
                max_cols = max(max_cols, len(row.cells))
            
            # Second pass: build table data
            for row_idx, row in enumerate(table.rows):
                row_data = []
                
                for col_idx in range(max_cols):
                    if col_idx < len(row.cells):
                        cell = row.cells[col_idx]
                        cell_text = cell.text.strip()
                        
                        if cell_text:
                            # Escape XML characters
                            cell_text = cell_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            
                            try:
                                # Try creating paragraph with Unicode text
                                cell_para = Paragraph(cell_text, body_style)
                                row_data.append(cell_para)
                            except Exception as e:
                                logger.debug(f"Paragraph creation failed for cell [{row_idx},{col_idx}]: {e}")
                                # Fallback to plain text
                                row_data.append(str(cell_text))
                        else:
                            row_data.append('')
                    else:
                        row_data.append('')
                
                table_data.append(row_data)
            
            if table_data and any(any(cell for cell in row) for row in table_data):
                # Calculate column widths dynamically
                available_width = 7.5 * inch  # Letter width minus margins
                col_widths = [available_width / max_cols] * max_cols
                
                # Create table with data and column widths
                t = Table(table_data, colWidths=col_widths, repeatRows=1)
                
                # Apply table styling with better visibility
                table_style = [
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),  # Blue header
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('FONTNAME', (0, 0), (-1, -1), main_font),
                    ('FONTSIZE', (0, 0), (-1, 0), 11),
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('LEFTPADDING', (0, 0), (-1, -1), 8),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#4472C4')),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F2F2F2')]),
                ]
                
                t.setStyle(TableStyle(table_style))
                
                story.append(Spacer(1, 0.2 * inch))
                story.append(t)
                story.append(Spacer(1, 0.3 * inch))
                logger.info(f"✓ Added table {table_idx + 1} with {len(table_data)} rows and {max_cols} columns")
            else:
                logger.warning(f"Table {table_idx + 1} has no data")
                
        except Exception as e:
            logger.error(f"Could not add table {table_idx + 1}: {e}")
            # Fallback: add table as formatted text
            try:
                story.append(Spacer(1, 0.1 * inch))
                story.append(Paragraph(f"<b>Table {table_idx + 1}:</b>", heading_style))
                
                for row_idx, row in enumerate(table.rows):
                    row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        row_text = ' | '.join(row_texts)
                        row_text = row_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        try:
                            p = Paragraph(row_text, body_style)
                            story.append(p)
                            story.append(Spacer(1, 0.05 * inch))
                        except:
                            logger.warning(f"Could not add row {row_idx} as text")
                
                story.append(Spacer(1, 0.2 * inch))
            except Exception as e2:
                logger.error(f"Fallback text rendering also failed: {e2}")
    
    # Build the PDF
    if story:
        try:
            pdf.build(story)
            logger.info("✓ Successfully created PDF using reportlab fallback")
        except Exception as e:
            logger.error(f"Failed to build PDF: {e}")
            # Create minimal PDF
            pdf.build([Paragraph("Error: Could not render document content", body_style)])
    else:
        # If no content, create a blank PDF
        pdf.build([Paragraph("No content found in document", body_style)])
        logger.warning("No content found in document")
    
    return output_path
